"""
generate_report.py - Gerador de Relatórios Executivos PROFISSIONAL
Versão Enterprise com gráficos de criticidade por arquivo
MELHORIAS:
  - Donut chart de distribuição de severidade
  - Gráfico de barras: score de criticidade por arquivo
  - Gráfico de barras: volume exposto por arquivo
  - Mapa de calor multidimensional de risco
  - Gauge individual (velocímetro) para cada arquivo vulnerável
  - Suporte completo AWS S3 / GCS / Azure Blob Storage
"""

# ── Matplotlib (backend não-interativo ANTES de qualquer import plt) ──
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

# ── ReportLab ─────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    PageBreak, KeepTogether, Image as RLImage
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT

# ── Stdlib ────────────────────────────────────────────────────────────
from datetime import datetime
from pathlib import Path
from collections import Counter
import hashlib, json, io, tempfile, os

# ── python-docx ───────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.shared import RGBColor, Pt, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ══════════════════════════════════════════════════════════════════════
# CHART ENGINE — gera PNGs em memória (BytesIO) ou arquivo temporário
# ══════════════════════════════════════════════════════════════════════
class ChartEngine:
    """Gera todos os gráficos como imagens PNG prontas para embutir."""

    PALETTE = {
        'primary':   '#1e3a8a',
        'secondary': '#3b82f6',
        'critical':  '#dc2626',
        'high':      '#ea580c',
        'medium':    '#ca8a04',
        'low':       '#16a34a',
        'bg':        '#f8fafc',
        'gray':      '#6b7280',
        'white':     '#ffffff',
    }

    SEV_COLOR = {
        'CRITICAL': '#dc2626',
        'HIGH':     '#ea580c',
        'MEDIUM':   '#ca8a04',
        'LOW':      '#16a34a',
    }

    # Score base por severidade (0-100)
    SEV_BASE = {'CRITICAL': 90, 'HIGH': 68, 'MEDIUM': 42, 'LOW': 18}

    def __init__(self, tmp_dir: Path):
        self.tmp = tmp_dir
        self.tmp.mkdir(exist_ok=True)

    def _save(self, fig, name: str) -> str:
        """Salva figura como PNG e retorna o caminho."""
        path = str(self.tmp / f'{name}.png')
        fig.savefig(path, dpi=150, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
        plt.close(fig)
        return path

    def _file_score(self, f: dict, max_size: float) -> float:
        """Calcula score de criticidade de um arquivo (0-100)."""
        base  = self.SEV_BASE.get(f.get('severity', 'LOW'), 18)
        size  = float(f.get('size', 0))
        bonus = (size / max_size * 12) if max_size > 0 else 0
        return min(100.0, base + bonus)

    # ── 1. Donut — distribuição de severidade ─────────────────────────
    def severity_donut(self, severity_dist: dict, total: int) -> str:
        P = self.PALETTE
        labels  = ['CRÍTICO', 'ALTO', 'MÉDIO', 'BAIXO']
        keys    = ['critical', 'high', 'medium', 'low']
        pal     = [P['critical'], P['high'], P['medium'], P['low']]
        values  = [severity_dist.get(k, 0) for k in keys]
        explode = [0.04, 0.04, 0.04, 0.01]

        nz_vals = [v for v in values if v > 0]
        nz_cols = [c for c, v in zip(pal, values) if v > 0]
        nz_exp  = [e for e, v in zip(explode, values) if v > 0]

        fig, ax = plt.subplots(figsize=(7, 4.8), facecolor='white')

        if sum(nz_vals) > 0:
            _, _, autotexts = ax.pie(
                nz_vals, colors=nz_cols, explode=nz_exp,
                autopct=lambda p: f'{p:.1f}%',
                startangle=140, pctdistance=0.75,
                wedgeprops={'linewidth': 2, 'edgecolor': 'white', 'antialiased': True}
            )
            for at in autotexts:
                at.set_fontsize(11); at.set_fontweight('bold'); at.set_color('white')

        circle = plt.Circle((0, 0), 0.55, color='white')
        ax.add_patch(circle)
        ax.text(0,  0.08, f'{total:,}',   ha='center', va='center',
                fontsize=20, fontweight='bold', color=P['primary'])
        ax.text(0, -0.16, 'arquivos', ha='center', va='center',
                fontsize=10, color=P['gray'])

        patches = [mpatches.Patch(color=pal[i], label=f'{labels[i]}: {values[i]:,}')
                   for i in range(4)]
        ax.legend(handles=patches, loc='lower center',
                  bbox_to_anchor=(0.5, -0.12), ncol=4, fontsize=9, frameon=False)

        ax.set_title('Distribuição de Severidade', fontsize=13,
                     fontweight='bold', color=P['primary'], pad=12)
        ax.axis('equal')
        fig.tight_layout()
        return self._save(fig, 'donut_severity')

    # ── 2. Barras — score de criticidade por arquivo ──────────────────
    def risk_per_file(self, files: list, title: str = '', fname: str = 'risk_files') -> str:
        if not files:
            return None
        P = self.PALETTE
        max_size = max(float(f.get('size', 0)) for f in files) or 1

        names  = [f.get('key', f.get('name', ''))[-30:] for f in files]
        scores = [self._file_score(f, max_size) for f in files]
        cols   = [self.SEV_COLOR.get(f.get('severity', 'LOW'), P['low']) for f in files]

        n = len(files)
        fig, ax = plt.subplots(figsize=(9.5, max(3.5, n * 0.44 + 1.2)), facecolor='white')
        ax.set_facecolor('#f9fafb')

        y = np.arange(n)
        bars = ax.barh(y, scores, color=cols, height=0.58,
                       edgecolor='white', linewidth=0.8, alpha=0.92)
        ax.grid(axis='x', color='white', linewidth=1.4)
        ax.set_axisbelow(True)

        for bar, score, f in zip(bars, scores, files):
            sev = f.get('severity', 'LOW')
            ax.text(bar.get_width() + 0.6,
                    bar.get_y() + bar.get_height() / 2,
                    f'{sev}  {score:.0f} pts',
                    va='center', ha='left', fontsize=8,
                    fontweight='bold', color=self.SEV_COLOR.get(sev, P['low']))

        # Zonas de fundo
        for x0, x1, c in [(0, 25, P['low']), (25, 50, P['medium']),
                           (50, 75, P['high']), (75, 115, P['critical'])]:
            ax.axvspan(x0, x1, alpha=0.04, color=c)
        for x, lbl, c in [(12.5,'BAIXO', P['low']), (37.5,'MÉDIO', P['medium']),
                           (62.5,'ALTO',  P['high']), (95,'CRÍTICO', P['critical'])]:
            ax.text(x, -0.8, lbl, ha='center', fontsize=7,
                    color=c, fontweight='bold', style='italic')

        ax.set_yticks(y)
        ax.set_yticklabels(names, fontsize=8.5, color='#374151')
        ax.set_xlim(0, 118)
        ax.set_xlabel('Score de Criticidade', fontsize=9, color=P['gray'])
        ax.set_title(f'Score de Criticidade por Arquivo{" — " + title if title else ""}',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=10)
        for sp in ['top', 'right']:
            ax.spines[sp].set_visible(False)
        ax.spines['left'].set_color('#e5e7eb')
        ax.spines['bottom'].set_color('#e5e7eb')
        ax.invert_yaxis()
        fig.tight_layout()
        return self._save(fig, fname)

    # ── 3. Barras — tamanho por arquivo ──────────────────────────────
    def size_per_file(self, files: list) -> str:
        if not files:
            return None
        P = self.PALETTE

        def fmt_size(b):
            b = float(b)
            for u in ['B', 'KB', 'MB', 'GB']:
                if b < 1024: return b, u
                b /= 1024
            return b, 'TB'

        names = [f.get('key', f.get('name', ''))[-30:] for f in files]
        sizes_raw = [float(f.get('size', 0)) for f in files]
        max_s = max(sizes_raw) or 1
        sizes_mb = [s / (1024*1024) for s in sizes_raw]
        cols = [self.SEV_COLOR.get(f.get('severity', 'LOW'), P['low']) for f in files]

        n = len(files)
        fig, ax = plt.subplots(figsize=(9.5, max(3.5, n * 0.44 + 1.2)), facecolor='white')
        ax.set_facecolor('#f9fafb')

        y = np.arange(n)
        bars = ax.barh(y, sizes_mb, color=cols, height=0.58,
                       edgecolor='white', linewidth=0.8, alpha=0.88)
        ax.grid(axis='x', color='white', linewidth=1.4)
        ax.set_axisbelow(True)

        for bar, raw in zip(bars, sizes_raw):
            val, unit = fmt_size(raw)
            ax.text(bar.get_width() + max(sizes_mb) * 0.01,
                    bar.get_y() + bar.get_height() / 2,
                    f'{val:.2f} {unit}', va='center', ha='left',
                    fontsize=8.5, color='#374151')

        ax.set_yticks(y)
        ax.set_yticklabels(names, fontsize=8.5, color='#374151')
        ax.set_xlabel('Tamanho (MB)', fontsize=9, color=P['gray'])
        ax.set_title('Volume de Dados Expostos por Arquivo',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=10)
        for sp in ['top', 'right']:
            ax.spines[sp].set_visible(False)
        ax.invert_yaxis()
        fig.tight_layout()
        return self._save(fig, 'size_per_file')

    # ── 4. Mapa de calor multidimensional ────────────────────────────
    def risk_heatmap(self, files: list) -> str:
        if not files:
            return None
        P   = self.PALETTE
        n   = min(len(files), 15)
        sel = files[:n]
        max_size = max(float(f.get('size', 0)) for f in sel) or 1

        names = [f.get('key', f.get('name', ''))[-22:] for f in sel]
        dims  = ['Exposição\nPública', 'Volume\nDados', 'Tipo\nArquivo',
                 'Risco\nPrivacidade', 'Score\nFinal']

        def row_scores(f):
            sev  = f.get('severity', 'LOW')
            size = float(f.get('size', 0))
            ext  = Path(f.get('key', f.get('name', ''))).suffix.lower()
            # Exposição: crítico=95, alto=75, médio=50, baixo=30
            exp  = {'CRITICAL':95,'HIGH':75,'MEDIUM':50,'LOW':30}.get(sev, 30)
            vol  = min(100, size / max_size * 100)
            typ  = 65 if ext in ['.sql','.db','.env','.key','.pem'] else \
                   55 if ext in ['.zip','.tar','.gz','.bak'] else \
                   50 if ext in ['.js','.py','.php'] else 40
            priv = 70 if any(x in f.get('key', f.get('name','')).lower()
                             for x in ['user','customer','person','cpf','password','secret']) else 45
            final = min(100, self._file_score(f, max_size))
            return [exp, vol, typ, priv, final]

        matrix = np.array([row_scores(f) for f in sel])

        fig, ax = plt.subplots(figsize=(10.5, max(4, n * 0.48 + 1.8)), facecolor='white')
        im = ax.imshow(matrix, cmap=plt.cm.RdYlGn_r, vmin=0, vmax=100, aspect='auto')

        for i in range(n):
            for j in range(len(dims)):
                val = matrix[i, j]
                tc  = 'white' if val > 65 or val < 20 else '#222222'
                ax.text(j, i, f'{val:.0f}', ha='center', va='center',
                        fontsize=8.5, fontweight='bold', color=tc)

        ax.set_xticks(range(len(dims)))
        ax.set_xticklabels(dims, fontsize=9, fontweight='bold', color=P['primary'])
        ax.set_yticks(range(n))
        ax.set_yticklabels(names, fontsize=8.5, color='#374151')
        ax.set_title('Mapa de Calor de Risco por Arquivo',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=12)

        cbar = fig.colorbar(im, ax=ax, fraction=0.025, pad=0.02)
        cbar.set_label('Nível de Risco (0–100)', fontsize=8, color=P['gray'])
        cbar.ax.tick_params(labelsize=7)
        fig.tight_layout()
        return self._save(fig, 'heatmap')

    # ── 5. Pizza — distribuição por tipo de arquivo ───────────────────
    def type_pie(self, files: list) -> str:
        if not files:
            return None
        P = self.PALETTE
        counts = Counter(
            Path(f.get('key', f.get('name', ''))).suffix.lower() or '(sem ext)'
            for f in files
        )
        # Agrupar raridades
        top = counts.most_common(6)
        others = sum(v for _, v in counts.most_common()[6:])
        if others:
            top.append(('outros', others))

        labels = [k for k, _ in top]
        values = [v for _, v in top]
        pal = [P['secondary'], P['primary'], P['critical'],
               P['medium'], P['low'], P['high'], P['gray']]

        fig, ax = plt.subplots(figsize=(6, 4.5), facecolor='white')
        wedges, texts, autotexts = ax.pie(
            values, labels=labels, colors=pal[:len(labels)],
            autopct='%1.1f%%', startangle=90,
            wedgeprops={'linewidth': 2, 'edgecolor': 'white'},
            pctdistance=0.78
        )
        for at in autotexts:
            at.set_fontsize(10); at.set_fontweight('bold'); at.set_color('white')
        for t in texts:
            t.set_fontsize(9); t.set_color('#374151')

        ax.set_title('Distribuição por Tipo de Arquivo',
                     fontsize=12, fontweight='bold', color=P['primary'], pad=12)
        fig.tight_layout()
        return self._save(fig, 'type_pie')

    # ── 6. Gauge individual por arquivo ──────────────────────────────
    def individual_gauge(self, f: dict, idx: int, max_size: float) -> str:
        score = self._file_score(f, max_size)
        sev   = f.get('severity', 'LOW')
        color = self.SEV_COLOR.get(sev, self.PALETTE['low'])
        name  = f.get('key', f.get('name', ''))

        fig = plt.figure(figsize=(5.5, 3.4), facecolor='white')
        ax  = fig.add_subplot(111, projection='polar')
        ax.set_facecolor('white')

        # Zonas coloridas em semicírculo (π → 0)
        zones = [
            (np.pi,       np.pi*0.75, self.PALETTE['low']),
            (np.pi*0.75,  np.pi*0.50, self.PALETTE['medium']),
            (np.pi*0.50,  np.pi*0.25, self.PALETTE['high']),
            (np.pi*0.25,  0,          self.PALETTE['critical']),
        ]
        for t0, t1, zcol in zones:
            th = np.linspace(t0, t1, 60)
            ax.fill_between(th, 0.62, 0.98, alpha=0.20, color=zcol)
            ax.plot(th, np.full(60, 0.98), color=zcol, linewidth=4.5,
                    solid_capstyle='butt')

        # Agulha
        needle = np.pi - (score / 100.0) * np.pi
        ax.annotate('', xy=(needle, 0.82), xytext=(0, 0),
                    arrowprops=dict(arrowstyle='->', color=color,
                                   lw=3.5, mutation_scale=20))

        ax.set_ylim(0, 1.1)
        ax.axis('off')

        # Score e severidade no centro
        ax.text(0,  0.10, f'{score:.0f}', ha='center', va='center',
                fontsize=22, fontweight='bold', color=color,
                transform=ax.transData)
        ax.text(0, -0.20, sev, ha='center', va='center',
                fontsize=10, fontweight='bold', color=color,
                transform=ax.transData)

        # Rótulos das zonas
        for ang, lbl, c in [
            (np.pi*0.875, 'LOW',  self.PALETTE['low']),
            (np.pi*0.625, 'MED',  self.PALETTE['medium']),
            (np.pi*0.375, 'HIGH', self.PALETTE['high']),
            (np.pi*0.125, 'CRIT', self.PALETTE['critical']),
        ]:
            ax.text(ang, 1.12, lbl, ha='center', va='center',
                    fontsize=7, color=c, fontweight='bold')

        short = (name[-22:] if len(name) > 22 else name)
        ax.set_title(short, fontsize=9, fontweight='bold',
                     color='#374151', pad=2, y=0.08)
        fig.tight_layout(pad=0.4)
        return self._save(fig, f'gauge_{idx:03d}')

    # ── Gerar todos os gauges ─────────────────────────────────────────
    def all_gauges(self, files: list) -> list:
        if not files:
            return []
        max_size = max(float(f.get('size', 0)) for f in files) or 1
        paths = []
        for i, f in enumerate(files):
            sev = f.get('severity', 'LOW')
            if sev in ('CRITICAL', 'HIGH', 'MEDIUM', 'LOW'):
                paths.append(self.individual_gauge(f, i, max_size))
                print(f"  📊 Gauge {i+1}/{len(files)}: {f.get('key', f.get('name',''))[-35:]}")
        return paths


# ══════════════════════════════════════════════════════════════════════
# MAIN GENERATOR CLASS
# ══════════════════════════════════════════════════════════════════════
class ProfessionalReportGenerator:
    """Gerador de relatórios executivos com gráficos de criticidade."""

    # ── Cores ReportLab ───────────────────────────────────────────────
    COLOR_PRIMARY   = colors.HexColor('#1e3a8a')
    COLOR_SECONDARY = colors.HexColor('#3b82f6')
    COLOR_ACCENT    = colors.HexColor('#60a5fa')
    COLOR_CRITICAL  = colors.HexColor('#dc2626')
    COLOR_HIGH      = colors.HexColor('#ea580c')
    COLOR_MEDIUM    = colors.HexColor('#ca8a04')
    COLOR_LOW       = colors.HexColor('#16a34a')
    COLOR_HEADER    = colors.HexColor('#0f172a')
    COLOR_BG_LIGHT  = colors.HexColor('#f8fafc')

    PROVIDER_NAMES = {
        'AWS_S3':      'AWS S3',
        'GCS':         'Google Cloud Storage',
        'AZURE':       'Azure Blob Storage',
        'AZURE_BLOB':  'Azure Blob Storage',
        'UNIVERSAL':   'Multi-Cloud',
    }

    SENSITIVE_EXTENSIONS = {
        '.env', '.pem', '.key', '.crt', '.cer', '.p12', '.pfx',
        '.sql', '.db', '.sqlite', '.mdb', '.bak', '.dump',
        '.yaml', '.yml', '.ini', '.conf', '.cfg',
        '.zip', '.tar', '.gz', '.rar', '.7z',
        '.py', '.php', '.rb', '.java', '.go',
        '.json', '.xml', '.csv',
    }

    SEV_BASE = {'CRITICAL': 90, 'HIGH': 68, 'MEDIUM': 42, 'LOW': 18}

    def __init__(self, scan_data: dict, client_info: dict = None):
        print("🔧 Inicializando gerador de relatórios...")
        print(f"📊 Dados recebidos: {list(scan_data.keys())}")

        self._validate_scan_data(scan_data)

        self.scan_data   = scan_data
        self.client_info = client_info or {}
        self.timestamp   = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.report_date = datetime.now()

        self.output_dir = Path("./reports_executive")
        self.output_dir.mkdir(exist_ok=True)

        # Diretório temporário para gráficos
        self.charts_dir = self.output_dir / "charts_tmp"
        self.charts_dir.mkdir(exist_ok=True)
        self.charts = ChartEngine(self.charts_dir)

        self.provider_name = self._get_provider_name()
        print(f"☁️  Provider: {self.provider_name}")

        # Análises de dados
        self.vulnerabilities    = self._safe(self._analyze_vulnerabilities,  [])
        self.recommendations    = self._safe(self._generate_recommendations,  [])
        self.risk_level         = self._safe(self._calculate_risk_level,      self._default_risk())
        self.compliance_status  = self._safe(self._assess_compliance,         [])
        self.extension_stats    = self._safe(self._analyze_extensions,        [])
        self.size_distribution  = self._safe(self._analyze_size_distribution, {})
        self.top_critical_files = self._safe(self._get_top_critical_files,    [])
        self.top_largest_files  = self._safe(self._get_top_largest_files,     [])
        self.previous_scan      = self._safe(self._load_previous_scan,        None)
        self._safe(self._save_current_scan_snapshot, None)

        print("📈 Gerando gráficos...")
        self._generate_all_charts()

    # ── Helpers de inicialização ───────────────────────────────────────
    def _safe(self, fn, default):
        try:
            return fn()
        except Exception as e:
            print(f"⚠️  {fn.__name__}: {e}")
            return default

    def _default_risk(self):
        return {'level':'DESCONHECIDO','score':0,'color':self.COLOR_MEDIUM,
                'action':'ANÁLISE NECESSÁRIA','critical_count':0,'high_count':0,'medium_count':0}

    def _validate_scan_data(self, sd):
        if not sd: raise ValueError("scan_data vazio")
        sd.setdefault('bucket', 'unknown-bucket')
        sd.setdefault('files', [])
        sd.setdefault('severity_distribution', self._calc_sev_dist(sd.get('files', [])))
        sd.setdefault('risk_score', 0)
        sd.setdefault('provider', 'UNIVERSAL')
        print("✅ Validação concluída")

    def _calc_sev_dist(self, files):
        dist = {'critical':0,'high':0,'medium':0,'low':0}
        for f in files:
            k = f.get('severity','LOW').lower()
            if k in dist: dist[k] += 1
        return dist

    def _get_provider_name(self):
        p = self.scan_data.get('provider','UNIVERSAL').upper()
        return self.PROVIDER_NAMES.get(p, p)

    # ── Análises ──────────────────────────────────────────────────────
    def _calculate_risk_level(self):
        sd    = self.scan_data.get('severity_distribution', {})
        score = self.scan_data.get('risk_score', 0)
        crit, high, med = sd.get('critical',0), sd.get('high',0), sd.get('medium',0)
        if crit > 0:    lvl, col, act = 'CRÍTICO', self.COLOR_CRITICAL, 'AÇÃO IMEDIATA REQUERIDA'
        elif high > 5:  lvl, col, act = 'ALTO',    self.COLOR_HIGH,     'AÇÃO URGENTE REQUERIDA'
        elif high > 0 or med > 10:
                        lvl, col, act = 'MÉDIO',   self.COLOR_MEDIUM,   'AÇÃO NECESSÁRIA'
        else:           lvl, col, act = 'BAIXO',   self.COLOR_LOW,      'MONITORAMENTO RECOMENDADO'
        return {'level':lvl,'score':score,'color':col,'action':act,
                'critical_count':crit,'high_count':high,'medium_count':med}

    def _assess_compliance(self):
        sd   = self.scan_data.get('severity_distribution', {})
        crit = sd.get('critical', 0)
        high = sd.get('high', 0)
        out  = []
        if crit > 0 or high > 0:
            out.append({'name':'LGPD / GDPR','status':'❌ NÃO CONFORME','issues':'Dados pessoais potencialmente expostos'})
        else:
            out.append({'name':'LGPD / GDPR','status':'⚠️ REVISAR','issues':'Verificar classificação de dados'})
        out.append({'name':'ISO 27001',
                    'status': '❌ NÃO CONFORME' if crit > 0 else '⚠️ PARCIAL',
                    'issues': 'Controles de acesso inadequados' if crit > 0 else 'Revisar políticas'})
        if any('card' in f.get('key','').lower() or 'payment' in f.get('key','').lower()
               for f in self.scan_data.get('files',[])):
            out.append({'name':'PCI DSS','status':'❌ CRÍTICO','issues':'Dados de pagamento expostos'})
        return out

    def _analyze_vulnerabilities(self):
        categories = {
            'credentials': {'name':'Credenciais Expostas',    'icon':'[KEY]','items':[]},
            'databases':   {'name':'Bancos de Dados',          'icon':'[DB]', 'items':[]},
            'config':      {'name':'Arq. de Configuração',     'icon':'[CFG]','items':[]},
            'backups':     {'name':'Backups',                  'icon':'[BAK]','items':[]},
            'source_code': {'name':'Código Fonte',             'icon':'[SRC]','items':[]},
            'keys':        {'name':'Chaves Criptográficas',    'icon':'[SEC]','items':[]},
            'pii':         {'name':'Dados Pessoais (PII)',     'icon':'[PII]','items':[]},
        }
        for f in self.scan_data.get('files', []):
            key = f.get('key', f.get('name', '')).lower()
            sev = f.get('severity', 'LOW')
            if sev not in ('CRITICAL', 'HIGH'): continue
            item = {'file': f.get('key', f.get('name','unknown')),
                    'size': f.get('size', 0), 'severity': sev,
                    'reason': f.get('reason',''), 'last_modified': f.get('last_modified','N/A')}
            if any(x in key for x in ['password','credential','secret','token','api_key','apikey']):
                categories['credentials']['items'].append(item)
            elif any(x in key for x in ['.sql','.db','.sqlite','database','dump','.mdb']):
                categories['databases']['items'].append(item)
            elif any(x in key for x in ['.env','config','.ini','.yaml','.yml','.conf','settings']):
                categories['config']['items'].append(item)
            elif any(x in key for x in ['backup','.bak','.old','.zip','.tar','.gz','.rar']):
                categories['backups']['items'].append(item)
            elif any(x in key for x in ['.pem','.key','.crt','.cer','.p12','.pfx','private','rsa']):
                categories['keys']['items'].append(item)
            elif any(x in key for x in ['cpf','rg','passport','social','personal','customer']):
                categories['pii']['items'].append(item)
            elif any(x in key for x in ['.py','.java','.js','.php','.rb','.go','.cpp']):
                categories['source_code']['items'].append(item)
        return [{'name':v['name'],'icon':v['icon'],'count':len(v['items']),'items':v['items'][:20]}
                for v in categories.values() if v['items']]

    def _generate_recommendations(self):
        sd   = self.scan_data.get('severity_distribution', {})
        prov = self.scan_data.get('provider','UNIVERSAL').upper()
        crit = sd.get('critical', 0)
        high = sd.get('high', 0)
        recs = []

        if crit > 0 or high > 0:
            actions_map = {
                'AWS_S3': ["Bloquear acesso público via 'Block Public Access'",
                           "Revisar bucket policies e ACLs no console AWS",
                           "Implementar bucket encryption (AES-256 ou KMS)",
                           "Habilitar versioning para recovery",
                           "Configurar logging para auditoria"],
                'GCS':    ["Remover 'allUsers' e 'allAuthenticatedUsers' das permissões",
                           "Ativar 'Uniform bucket-level access'",
                           "Implementar IAM Conditions",
                           "Habilitar 'Object versioning'",
                           "Configurar 'Audit logs'"],
            }
            if 'AZURE' in prov:
                actions = ["Revisar IAM no Azure Portal",
                           "Configurar 'Private endpoints'",
                           "Habilitar 'SAS' com expiração",
                           "Implementar Azure AD",
                           "Configurar 'Network rules'"]
            else:
                actions = actions_map.get(prov, ["Revisar políticas de acesso",
                                                  "Remover permissões públicas",
                                                  "Implementar autenticação forte",
                                                  "Habilitar criptografia em repouso",
                                                  "Configurar logs de auditoria"])
            recs.append({'priority':'CRÍTICA','title':'Restringir Acesso Público ao Storage',
                         'description':f'{crit+high} arquivo(s) com exposição pública detectados.',
                         'actions':actions,'timeline':'0-24 horas',
                         'responsible':'Equipe de Segurança / DevOps'})

        recs.append({'priority':'ALTA','title':'Implementar Criptografia de Dados',
                     'description':'Garantir criptografia em repouso e em trânsito.',
                     'actions':["Habilitar criptografia server-side",
                                "Implementar HTTPS/TLS para transferências",
                                "Rotacionar chaves regularmente",
                                "Usar CMK quando possível",
                                "Documentar gestão de chaves"],
                     'timeline':'7-14 dias','responsible':'Equipe de Segurança'})

        recs.append({'priority':'MÉDIA','title':'Implementar Monitoramento Contínuo',
                     'description':'Detectar acessos suspeitos e mudanças de configuração.',
                     'actions':["Configurar alertas para acessos anômalos",
                                "Implementar SIEM para análise de logs",
                                "Criar dashboards de segurança",
                                "Estabelecer processo de resposta a incidentes",
                                "Realizar auditorias regulares"],
                     'timeline':'14-30 dias','responsible':'Equipe de SecOps'})

        recs.append({'priority':'MÉDIA','title':'Estabelecer Políticas de Governança',
                     'description':'Classificação e proteção de dados com políticas claras.',
                     'actions':["Classificar dados por sensibilidade",
                                "Definir políticas de retenção",
                                "Implementar DLP",
                                "Treinar equipes",
                                "Revisões trimestrais de acessos"],
                     'timeline':'30-60 dias','responsible':'CISO / Compliance'})
        return recs

    def _analyze_extensions(self):
        ext_map = {}
        for f in self.scan_data.get('files', []):
            ext = Path(f.get('key', f.get('name',''))).suffix.lower() or '(sem extensão)'
            sev = f.get('severity','LOW').lower()
            if ext not in ext_map:
                ext_map[ext] = {'total':0,'critical':0,'high':0,'medium':0,'low':0,
                                'sensitive': ext in self.SENSITIVE_EXTENSIONS}
            ext_map[ext]['total'] += 1
            if sev in ext_map[ext]: ext_map[ext][sev] += 1
        result = [{'ext':k,**v} for k,v in ext_map.items()]
        result.sort(key=lambda x: (x['critical']+x['high'], x['total']), reverse=True)
        return result[:25]

    def _analyze_size_distribution(self):
        stats = {s:{'count':0,'total_bytes':0} for s in ['critical','high','medium','low']}
        for f in self.scan_data.get('files', []):
            sev  = f.get('severity','LOW').lower()
            size = int(f.get('size', 0))
            if sev in stats:
                stats[sev]['count']       += 1
                stats[sev]['total_bytes'] += size
        for k in stats:
            c = stats[k]['count']
            stats[k]['avg_bytes'] = stats[k]['total_bytes'] // c if c else 0
        stats['_total_bytes'] = sum(v['total_bytes'] for v in stats.values())
        return stats

    def _get_top_critical_files(self, n=15):
        order = {'CRITICAL':0,'HIGH':1,'MEDIUM':2,'LOW':3}
        files = sorted(self.scan_data.get('files',[]),
                       key=lambda f: (order.get(f.get('severity','LOW'),4),
                                      -int(f.get('size',0))))
        return [{'file': f.get('key',f.get('name','unknown')),
                 'severity': f.get('severity','LOW'),
                 'size': int(f.get('size',0)),
                 'reason': f.get('reason','Exposição pública'),
                 'last_modified': f.get('last_modified','N/A')}
                for f in files[:n]
                if f.get('severity','LOW') in ('CRITICAL','HIGH','MEDIUM')]

    def _get_top_largest_files(self, n=10):
        files = sorted(self.scan_data.get('files',[]),
                       key=lambda f: int(f.get('size',0)), reverse=True)
        return [{'file': f.get('key',f.get('name','unknown')),
                 'severity': f.get('severity','LOW'),
                 'size': int(f.get('size',0))} for f in files[:n]]

    def _load_previous_scan(self):
        snap = self.output_dir / f"snapshot_{self.scan_data.get('bucket','').replace('.','_')}.json"
        if snap.exists():
            with open(snap) as fh: return json.load(fh)
        return None

    def _save_current_scan_snapshot(self):
        snap = {'timestamp': self.timestamp,
                'severity_distribution': self.scan_data.get('severity_distribution',{}),
                'total_files': len(self.scan_data.get('files',[])),
                'risk_score':  self.scan_data.get('risk_score',0)}
        path = self.output_dir / f"snapshot_{self.scan_data.get('bucket','').replace('.','_')}.json"
        with open(path,'w') as fh: json.dump(snap, fh, indent=2)

    def _build_comparison(self):
        if not self.previous_scan: return None
        curr = self.scan_data.get('severity_distribution',{})
        prev = self.previous_scan.get('severity_distribution',{})
        return {'prev_timestamp': self.previous_scan.get('timestamp','N/A'),
                'prev_total':     self.previous_scan.get('total_files',0),
                'curr_total':     len(self.scan_data.get('files',[])),
                'delta_critical': curr.get('critical',0) - prev.get('critical',0),
                'delta_high':     curr.get('high',0)     - prev.get('high',0),
                'delta_medium':   curr.get('medium',0)   - prev.get('medium',0),
                'delta_score':    self.scan_data.get('risk_score',0) - self.previous_scan.get('risk_score',0)}

    # ── Gerar todos os gráficos ────────────────────────────────────────
    def _generate_all_charts(self):
        files    = self.scan_data.get('files', [])
        sd       = self.scan_data.get('severity_distribution', {})
        total    = len(files)
        vuln     = [f for f in files if f.get('severity','LOW') != 'LOW']
        # Incluir também LOW se não houver vulneráveis (para garantir gauges)
        gauge_files = vuln if vuln else files[:20]

        self.ch = {}
        self.ch['donut']    = self._safe(lambda: self.charts.severity_donut(sd, total), None)
        self.ch['risk_all'] = self._safe(lambda: self.charts.risk_per_file(files[:20], 'Top 20'), None)
        self.ch['risk_vuln']= self._safe(lambda: self.charts.risk_per_file(gauge_files, 'Vulneráveis', 'risk_vuln'), None) if vuln else self.ch['risk_all']
        self.ch['size']     = self._safe(lambda: self.charts.size_per_file(files[:20]), None)
        self.ch['heatmap']  = self._safe(lambda: self.charts.risk_heatmap(files[:15]), None)
        self.ch['type_pie'] = self._safe(lambda: self.charts.type_pie(files), None)
        self.ch['gauges']   = self._safe(lambda: self.charts.all_gauges(gauge_files[:20]), [])
        print(f"✅ {sum(1 for v in self.ch.values() if v)} gráficos gerados "
              f"({len(self.ch['gauges'])} gauges individuais)")

    # ── Helpers de formatação ──────────────────────────────────────────
    def _format_size(self, b):
        try: b = int(b)
        except: return "0 B"
        for u in ['B','KB','MB','GB']:
            if b < 1024: return f"{b:.1f} {u}"
            b /= 1024
        return f"{b:.1f} TB"

    def _calc_pct(self, v, t):
        return f"{v/t*100:.1f}" if t else "0.0"

    def _delta_str(self, v):
        return f"+{v}" if v > 0 else str(v)

    def _compute_hash(self, filepath):
        h = hashlib.sha256()
        with open(filepath,'rb') as fh:
            for chunk in iter(lambda: fh.read(8192), b''):
                h.update(chunk)
        return h.hexdigest()

    # ══════════════════════════════════════════════════════════════════
    # PDF GENERATION
    # ══════════════════════════════════════════════════════════════════
    def _create_header_footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFillColor(self.COLOR_PRIMARY)
        canvas.rect(0, A4[1]-0.6*inch, A4[0], 0.6*inch, fill=True, stroke=False)
        canvas.setFillColor(colors.white)
        canvas.setFont('Helvetica-Bold', 12)
        canvas.drawString(0.5*inch, A4[1]-0.35*inch, "RELATORIO DE SEGURANCA EXECUTIVO")
        canvas.setFont('Helvetica', 9)
        canvas.drawRightString(A4[0]-0.5*inch, A4[1]-0.35*inch,
                               self.report_date.strftime("%d/%m/%Y"))
        canvas.setFillColor(colors.grey)
        canvas.setFont('Helvetica', 8)
        canvas.drawString(0.5*inch, 0.4*inch,
                          f"Security Multicloud Scanner | {self.provider_name}")
        canvas.drawRightString(A4[0]-0.5*inch, 0.4*inch, f"Pagina {doc.page}")
        canvas.setStrokeColor(self.COLOR_PRIMARY)
        canvas.setLineWidth(2)
        canvas.line(0.5*inch, 0.6*inch, A4[0]-0.5*inch, 0.6*inch)
        canvas.restoreState()

    def generate_pdf(self) -> str:
        print("📄 Gerando PDF...")
        bucket   = self.scan_data.get('bucket','scan').replace('.','_').replace('/','_')
        filename = f"relatorio_{bucket}_{self.timestamp}.pdf"
        filepath = self.output_dir / filename

        doc   = SimpleDocTemplate(str(filepath), pagesize=A4,
                                  topMargin=0.8*inch, bottomMargin=0.8*inch,
                                  leftMargin=0.7*inch, rightMargin=0.7*inch)
        story = []
        styles = getSampleStyleSheet()

        H1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=16,
                             textColor=self.COLOR_PRIMARY, spaceAfter=12,
                             spaceBefore=12, fontName='Helvetica-Bold')
        H2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12,
                             textColor=self.COLOR_SECONDARY, spaceAfter=8,
                             spaceBefore=8, fontName='Helvetica-Bold')
        NOTE = ParagraphStyle('NOTE', fontSize=8, textColor=colors.grey,
                              alignment=TA_CENTER, fontName='Helvetica-Oblique')

        def add_chart(path, width=5.8*inch, caption=None):
            if path and Path(path).exists():
                story.append(Spacer(1, 0.1*inch))
                story.append(RLImage(path, width=width,
                                     height=width * 0.55))
                if caption:
                    story.append(Paragraph(caption, NOTE))
                story.append(Spacer(1, 0.15*inch))

        sd     = self.scan_data.get('severity_distribution', {})
        total  = len(self.scan_data.get('files', []))
        risk   = self.risk_level

        # ── CAPA ──────────────────────────────────────────────────────
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph("RELATORIO DE SEGURANCA EXECUTIVO",
                                ParagraphStyle('T', fontSize=26, alignment=TA_CENTER,
                                               textColor=self.COLOR_PRIMARY,
                                               fontName='Helvetica-Bold', spaceAfter=30, leading=32)))
        story.append(Paragraph(f"Auditoria de Storage Multicloud<br/>{self.provider_name}",
                                ParagraphStyle('S', fontSize=14, alignment=TA_CENTER,
                                               textColor=colors.grey, spaceAfter=40)))
        if self.client_info.get('name'):
            ct = Table([['Cliente:', self.client_info.get('name','')],
                        ['Contato:', self.client_info.get('contact','')],
                        ['Data:', self.report_date.strftime('%d/%m/%Y às %H:%M')],
                        ['Alvo:', self.scan_data.get('bucket','')]],
                       colWidths=[1.5*inch, 3.5*inch])
            ct.setStyle(TableStyle([
                ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),11),
                ('TEXTCOLOR',(0,0),(0,-1),self.COLOR_PRIMARY),
                ('PADDING',(0,0),(-1,-1),8),
            ]))
            story.append(ct)
        story.append(PageBreak())

        # ── 1. RESUMO EXECUTIVO ────────────────────────────────────────
        story.append(Paragraph("1. RESUMO EXECUTIVO", H1))
        # Risk box
        rb = Table([[
            Paragraph(f"<b>NÍVEL DE RISCO: {risk['level']}</b>",
                      ParagraphStyle('RL', fontSize=14, textColor=colors.white, alignment=TA_CENTER)),
            Paragraph(f"<b>{risk['score']}%</b>",
                      ParagraphStyle('RS', fontSize=20, textColor=colors.white,
                                     alignment=TA_CENTER, fontName='Helvetica-Bold'))
        ]], colWidths=[3*inch, 1.5*inch])
        rb.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,-1),risk['color']),
            ('PADDING',(0,0),(-1,-1),15),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ]))
        story.append(rb)
        story.append(Spacer(1, 0.2*inch))

        # Métricas
        mt = Table([
            ['Métrica','Valor','% do Total'],
            ['Provider', self.provider_name, ''],
            ['Bucket/Container', self.scan_data.get('bucket',''), ''],
            ['Total de Arquivos', str(total), '100%'],
            ['CRÍTICOS', str(sd.get('critical',0)),
             f"{self._calc_pct(sd.get('critical',0),total)}%"],
            ['ALTOS',    str(sd.get('high',0)),
             f"{self._calc_pct(sd.get('high',0),total)}%"],
            ['MÉDIOS',   str(sd.get('medium',0)),
             f"{self._calc_pct(sd.get('medium',0),total)}%"],
            ['BAIXOS',   str(sd.get('low',0)),
             f"{self._calc_pct(sd.get('low',0),total)}%"],
            ['Tamanho Total', self._format_size(self.size_distribution.get('_total_bytes',0)), ''],
        ], colWidths=[3*inch, 2*inch, 1.2*inch])
        mt.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0),self.COLOR_PRIMARY),
            ('TEXTCOLOR',(0,0),(-1,0),colors.white),
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,-1),10),
            ('PADDING',(0,0),(-1,-1),10),
            ('GRID',(0,0),(-1,-1),0.5,self.COLOR_PRIMARY),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,self.COLOR_BG_LIGHT]),
            ('TEXTCOLOR',(1,4),(2,4),self.COLOR_CRITICAL),
            ('TEXTCOLOR',(1,5),(2,5),self.COLOR_HIGH),
            ('TEXTCOLOR',(1,6),(2,6),self.COLOR_MEDIUM),
            ('TEXTCOLOR',(1,7),(2,7),self.COLOR_LOW),
            ('FONTNAME',(1,4),(2,7),'Helvetica-Bold'),
        ]))
        story.append(mt)
        story.append(Spacer(1,0.15*inch))
        story.append(Paragraph(f"<b>{risk['action']}</b>",
                                ParagraphStyle('ACT', fontSize=12, textColor=risk['color'],
                                               alignment=TA_CENTER, fontName='Helvetica-Bold')))
        story.append(PageBreak())

        # ── 2. DISTRIBUIÇÃO DE CRITICIDADE ────────────────────────────
        story.append(Paragraph("2. DISTRIBUIÇÃO DE CRITICIDADE", H1))
        story.append(Paragraph(
            "Distribuição dos arquivos por nível de severidade e tipo de conteúdo exposto.",
            styles['Normal']))

        # Donut + Tipo lado a lado
        if self.ch.get('donut') and self.ch.get('type_pie'):
            row = Table([[
                RLImage(self.ch['donut'],   width=2.9*inch, height=2.9*inch*0.7),
                RLImage(self.ch['type_pie'],width=2.9*inch, height=2.9*inch*0.7),
            ]], colWidths=[3.1*inch, 3.1*inch])
            row.setStyle(TableStyle([('ALIGN',(0,0),(-1,-1),'CENTER'),
                                     ('VALIGN',(0,0),(-1,-1),'MIDDLE')]))
            story.append(row)
            story.append(Paragraph(
                "Figura 1 — Severidade (esq.) | Tipo de Arquivo (dir.)", NOTE))
        elif self.ch.get('donut'):
            add_chart(self.ch['donut'], caption="Figura 1 — Distribuição de Severidade")

        story.append(PageBreak())

        # ── 3. SCORE DE CRITICIDADE POR ARQUIVO ───────────────────────
        story.append(Paragraph("3. SCORE DE CRITICIDADE POR ARQUIVO", H1))
        story.append(Paragraph(
            "Score calculado por arquivo: base de severidade + bônus proporcional ao tamanho. "
            "Zonas coloridas indicam faixas de risco.", styles['Normal']))
        add_chart(self.ch.get('risk_all'), width=5.8*inch,
                  caption="Figura 2 — Score de Criticidade por Arquivo")
        story.append(Paragraph("Volume de dados expostos por arquivo:", H2))
        add_chart(self.ch.get('size'), width=5.8*inch,
                  caption="Figura 3 — Volume de Dados Expostos por Arquivo")
        story.append(PageBreak())

        # ── 4. MAPA DE CALOR ──────────────────────────────────────────
        story.append(Paragraph("4. MAPA DE CALOR DE RISCO MULTIDIMENSIONAL", H1))
        story.append(Paragraph(
            "Cada arquivo é avaliado em 5 dimensões: Exposição Pública, Volume de Dados, "
            "Tipo de Arquivo, Risco de Privacidade e Score Final. Cores quentes = maior risco.",
            styles['Normal']))
        add_chart(self.ch.get('heatmap'), width=6.0*inch,
                  caption="Figura 4 — Mapa de Calor Multidimensional de Risco")
        story.append(PageBreak())

        # ── 5. GAUGES INDIVIDUAIS ──────────────────────────────────────
        if self.ch.get('gauges'):
            story.append(Paragraph("5. ANÁLISE INDIVIDUAL POR ARQUIVO (GAUGES)", H1))
            story.append(Paragraph(
                "Velocímetro de criticidade individual para cada arquivo vulnerável. "
                "O ponteiro indica o score calculado (0–100).", styles['Normal']))
            story.append(Spacer(1, 0.15*inch))

            # Grade 3 colunas de gauges
            COLS = 3
            gauges = self.ch['gauges']
            files  = self.scan_data.get('files', [])[:len(gauges)]
            for row_start in range(0, len(gauges), COLS):
                chunk = gauges[row_start:row_start+COLS]
                imgs  = [RLImage(p, width=1.95*inch, height=1.95*inch*0.62) for p in chunk]
                while len(imgs) < COLS:
                    imgs.append(Paragraph('', styles['Normal']))
                row_t = Table([imgs], colWidths=[2.1*inch]*COLS)
                row_t.setStyle(TableStyle([
                    ('ALIGN',(0,0),(-1,-1),'CENTER'),
                    ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                    ('BOTTOMPADDING',(0,0),(-1,-1),4),
                ]))
                story.append(row_t)
            story.append(Paragraph(
                f"Figura 5 — Gauges de criticidade individual ({len(gauges)} arquivos)", NOTE))
            story.append(PageBreak())
            next_sec = 6
        else:
            next_sec = 5

        # ── VULNERABILIDADES ──────────────────────────────────────────
        if self.vulnerabilities:
            story.append(Paragraph(f"{next_sec}. VULNERABILIDADES IDENTIFICADAS", H1))
            for vuln in self.vulnerabilities:
                story.append(Paragraph(
                    f"{vuln['icon']} {vuln['name']} ({vuln['count']} arquivos)", H2))
                vt = Table(
                    [['Arquivo','Sev.','Tamanho','Motivo']] +
                    [[item['file'][:45]+'...' if len(item['file'])>45 else item['file'],
                      item['severity'],
                      self._format_size(item['size']),
                      item.get('reason','')[:30]] for item in vuln['items'][:10]],
                    colWidths=[2.8*inch, 0.7*inch, 0.8*inch, 1.8*inch])
                vt.setStyle(TableStyle([
                    ('BACKGROUND',(0,0),(-1,0),self.COLOR_SECONDARY),
                    ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                    ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                    ('FONTSIZE',(0,0),(-1,-1),8),
                    ('PADDING',(0,0),(-1,-1),6),
                    ('GRID',(0,0),(-1,-1),0.5,colors.grey),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,self.COLOR_BG_LIGHT]),
                ]))
                story.append(vt)
                story.append(Spacer(1, 0.15*inch))
            story.append(PageBreak())
            next_sec += 1

        # ── COMPARATIVO ───────────────────────────────────────────────
        comp = self._build_comparison()
        if comp:
            story.append(Paragraph(f"{next_sec}. COMPARATIVO COM SCAN ANTERIOR", H1))
            story.append(Paragraph(
                f"Scan anterior realizado em: <b>{comp['prev_timestamp']}</b>",
                styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            ct_data = [['Métrica','Anterior','Atual','Variação']]
            for lbl, prev_val, curr_val, delta in [
                ('Total', comp['prev_total'], comp['curr_total'],
                 comp['curr_total']-comp['prev_total']),
                ('CRÍTICOS',
                 self.previous_scan['severity_distribution'].get('critical',0),
                 self.scan_data['severity_distribution'].get('critical',0),
                 comp['delta_critical']),
                ('ALTOS',
                 self.previous_scan['severity_distribution'].get('high',0),
                 self.scan_data['severity_distribution'].get('high',0),
                 comp['delta_high']),
                ('Risk Score', self.previous_scan.get('risk_score',0),
                 self.scan_data.get('risk_score',0), comp['delta_score']),
            ]:
                ct_data.append([lbl, str(prev_val), str(curr_val), self._delta_str(delta)])

            comp_t = Table(ct_data, colWidths=[2.5*inch,1.2*inch,1.2*inch,1.2*inch])
            comp_t.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,0),self.COLOR_PRIMARY),
                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),10),
                ('PADDING',(0,0),(-1,-1),10),
                ('GRID',(0,0),(-1,-1),0.5,self.COLOR_PRIMARY),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,self.COLOR_BG_LIGHT]),
                ('ALIGN',(1,0),(-1,-1),'CENTER'),
            ]))
            story.append(comp_t)
            story.append(PageBreak())
            next_sec += 1

        # ── RECOMENDAÇÕES ─────────────────────────────────────────────
        story.append(Paragraph(f"{next_sec}. RECOMENDAÇÕES PRIORITÁRIAS", H1))
        effort_map = {'CRÍTICA':'Baixo','ALTA':'Médio','MÉDIA':'Alto','BAIXA':'Baixo'}
        impact_map = {'CRÍTICA':'Máximo','ALTA':'Alto','MÉDIA':'Médio','BAIXA':'Baixo'}
        pcols = {'CRÍTICA':self.COLOR_CRITICAL,'ALTA':self.COLOR_HIGH,
                 'MÉDIA':self.COLOR_MEDIUM,'BAIXA':self.COLOR_LOW}
        for rec in self.recommendations:
            story.append(Paragraph(
                f"[{rec['priority']}] {rec['title']}",
                ParagraphStyle('RT', fontSize=11, fontName='Helvetica-Bold',
                               textColor=pcols.get(rec['priority'],self.COLOR_MEDIUM), spaceAfter=5)))
            story.append(Paragraph(rec['description'], styles['Normal']))
            meta = Table([[
                'Prazo', rec.get('timeline','A definir'),
                'Responsável', rec.get('responsible','A definir'),
                'Esforço', effort_map.get(rec['priority'],'Médio'),
                'Impacto', impact_map.get(rec['priority'],'Alto'),
            ]], colWidths=[0.8*inch,1.3*inch,1.0*inch,1.7*inch,0.8*inch,0.9*inch,0.7*inch,0.8*inch])
            meta.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,-1),self.COLOR_BG_LIGHT),
                ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
                ('FONTNAME',(2,0),(2,-1),'Helvetica-Bold'),
                ('FONTNAME',(4,0),(4,-1),'Helvetica-Bold'),
                ('FONTNAME',(6,0),(6,-1),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),8),
                ('PADDING',(0,0),(-1,-1),5),
                ('GRID',(0,0),(-1,-1),0.3,colors.grey),
                ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ]))
            story.append(Spacer(1,0.05*inch))
            story.append(meta)
            story.append(Spacer(1,0.05*inch))
            story.append(Paragraph("<b>Ações:</b>", styles['Normal']))
            for a in rec['actions'][:5]:
                story.append(Paragraph(f"• {a}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        story.append(PageBreak())
        next_sec += 1

        # ── COMPLIANCE ────────────────────────────────────────────────
        if self.compliance_status:
            story.append(Paragraph(f"{next_sec}. STATUS DE CONFORMIDADE", H1))
            comp_t2 = Table(
                [['Framework','Status','Observações']] +
                [[fw['name'],fw['status'],fw['issues']] for fw in self.compliance_status],
                colWidths=[1.5*inch,1.8*inch,2.8*inch])
            comp_t2.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,0),self.COLOR_PRIMARY),
                ('TEXTCOLOR',(0,0),(-1,0),colors.white),
                ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),9),
                ('PADDING',(0,0),(-1,-1),8),
                ('GRID',(0,0),(-1,-1),0.5,self.COLOR_PRIMARY),
                ('VALIGN',(0,0),(-1,-1),'TOP'),
            ]))
            story.append(comp_t2)
            story.append(Spacer(1,0.3*inch))
            next_sec += 1

        # ── CONCLUSÃO ─────────────────────────────────────────────────
        story.append(Paragraph(f"{next_sec}. CONCLUSÕES E PRÓXIMOS PASSOS", H1))
        story.append(Paragraph(
            f"Este relatório identificou <b>{risk['critical_count']} vulnerabilidades críticas</b> e "
            f"<b>{risk['high_count']} de alto risco</b> no ambiente {self.provider_name}. "
            "É fundamental executar as ações prioritárias nos prazos estabelecidos.<br/><br/>"
            "<b>Próximos Passos:</b><br/>"
            "1. Reunião de alinhamento com stakeholders (48h)<br/>"
            "2. Execução das ações críticas (0-7 dias)<br/>"
            "3. Controles preventivos (30-60 dias)<br/>"
            "4. Re-auditoria de segurança (90 dias)<br/>"
            "5. Programa de segurança contínua",
            styles['Normal']))
        story.append(Spacer(1,0.5*inch))
        story.append(Paragraph(
            f"<b>Security Multicloud Scanner</b><br/>{self.provider_name} | "
            "Relatório gerado automaticamente<br/>Este documento contém informações confidenciais",
            ParagraphStyle('SIG', fontSize=8, textColor=colors.grey, alignment=TA_CENTER)))

        print("🔨 Construindo PDF...")
        doc.build(story, onFirstPage=self._create_header_footer,
                  onLaterPages=self._create_header_footer)

        report_hash = self._compute_hash(str(filepath))
        (filepath.with_suffix('.sha256')).write_text(
            f"SHA-256: {report_hash}\nArquivo: {filename}\nGerado em: {self.report_date.isoformat()}\n")
        print(f"✅ PDF: {filepath}")
        return str(filepath)

    # ══════════════════════════════════════════════════════════════════
    # DOCX GENERATION
    # ══════════════════════════════════════════════════════════════════
    def _docx_set_bg(self, cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color.replace('#',''))
        tcPr.append(shd)

    def _docx_remove_borders(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        bdr  = OxmlElement('w:tcBorders')
        for side in ['top','bottom','left','right','insideH','insideV']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            bdr.append(el)
        tcPr.append(bdr)

    def _docx_heading(self, doc, text, level=1):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt({1:16,2:13,3:11}.get(level,11))
        run.font.color.rgb = RGBColor(0x1e,0x3a,0x8a) if level==1 else RGBColor(0x3b,0x82,0xf6)
        run.font.name = 'Arial'
        p.paragraph_format.space_before = Pt(14 if level==1 else 10)
        p.paragraph_format.space_after  = Pt(6)
        if level == 1:
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
            bot.set(qn('w:space'),'4');   bot.set(qn('w:color'),'3b82f6')
            pBdr.append(bot); pPr.append(pBdr)
        return p

    def _docx_body(self, doc, text, color='333333', bold=False, size=10, after=4):
        p   = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(after)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        run.font.size  = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name  = 'Arial'
        run.bold       = bold
        return p

    def _docx_add_chart(self, doc, path, width_in=6.0, caption=None):
        if not path or not Path(path).exists(): return
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run().add_picture(path, width=Inches(width_in))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            r = cap.add_run(caption)
            r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x6b,0x72,0x80)
            r.font.name = 'Arial'; r.italic = True

    def generate_docx(self) -> str:
        if not DOCX_AVAILABLE:
            print("⚠️  python-docx indisponível"); return None
        print("📝 Gerando DOCX...")
        bucket   = self.scan_data.get('bucket','scan').replace('.','_').replace('/','_')
        filename = f"relatorio_{bucket}_{self.timestamp}.docx"
        filepath = self.output_dir / filename

        doc = Document()
        for sec in doc.sections:
            sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
            sec.left_margin   = Inches(0.85);  sec.right_margin  = Inches(0.85)
            sec.top_margin    = Inches(0.9);   sec.bottom_margin = Inches(0.9)

        sd    = self.scan_data.get('severity_distribution', {})
        total = len(self.scan_data.get('files', []))
        risk  = self.risk_level

        # ── Capa ──────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        r = p.add_run('RELATÓRIO DE SEGURANÇA EXECUTIVO\nCOM ANÁLISE DE CRITICIDADE')
        r.font.size = Pt(24); r.bold = True
        r.font.color.rgb = RGBColor(0x1e,0x3a,0x8a); r.font.name = 'Arial'
        self._docx_body(doc, f'Auditoria de Storage Multicloud — {self.provider_name}',
                         color='6b7280', size=13).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        if self.client_info.get('name'):
            tbl = doc.add_table(rows=4, cols=2)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            rows_data = [('Cliente', self.client_info.get('name','')),
                         ('Contato', self.client_info.get('contact','')),
                         ('Data', self.report_date.strftime('%d/%m/%Y às %H:%M')),
                         ('Alvo', self.scan_data.get('bucket',''))]
            for i,(k,v) in enumerate(rows_data):
                self._docx_set_bg(tbl.rows[i].cells[0], 'e8eef8')
                tbl.rows[i].cells[0].width = Inches(1.4)
                tbl.rows[i].cells[1].width = Inches(3.0)
                r0 = tbl.rows[i].cells[0].paragraphs[0].add_run(k)
                r0.bold=True; r0.font.size=Pt(10); r0.font.name='Arial'
                r0.font.color.rgb=RGBColor(0x1e,0x3a,0x8a)
                r1 = tbl.rows[i].cells[1].paragraphs[0].add_run(v)
                r1.font.size=Pt(10); r1.font.name='Arial'

        doc.add_page_break()

        # ── 1. Resumo ─────────────────────────────────────────────────
        self._docx_heading(doc, '1. RESUMO EXECUTIVO')
        self._docx_body(doc, f"Nível de Risco: {risk['level']} ({risk['score']}%) — {risk['action']}",
                         bold=True, color='dc2626' if risk['level']=='CRÍTICO' else '1e3a8a')
        doc.add_paragraph()

        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        metrics = [('Arquivos', str(total), '1e3a8a'),
                   ('CRÍTICOS', str(sd.get('critical',0)), 'dc2626'),
                   ('ALTOS',    str(sd.get('high',0)),    'ea580c'),
                   ('Tamanho',  self._format_size(self.size_distribution.get('_total_bytes',0)), '1e3a8a')]
        for i,(lbl,val,col) in enumerate(metrics):
            self._docx_set_bg(tbl.rows[0].cells[i], '1e3a8a')
            r0 = tbl.rows[0].cells[i].paragraphs[0].add_run(lbl)
            r0.bold=True; r0.font.size=Pt(9); r0.font.name='Arial'
            r0.font.color.rgb=RGBColor(255,255,255)
            tbl.rows[0].cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            r1 = tbl.rows[1].cells[i].paragraphs[0].add_run(val)
            r1.bold=True; r1.font.size=Pt(16); r1.font.name='Arial'
            r1.font.color.rgb=RGBColor.from_string(col)
            tbl.rows[1].cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ── 2. Distribuição de criticidade ───────────────────────────
        self._docx_heading(doc, '2. DISTRIBUIÇÃO DE CRITICIDADE')
        self._docx_body(doc, 'Gráfico de distribuição dos arquivos por nível de severidade e tipo de conteúdo.')
        self._docx_add_chart(doc, self.ch.get('donut'), width_in=5.5,
                              caption='Figura 1 — Distribuição de Severidade')
        if self.ch.get('type_pie'):
            self._docx_add_chart(doc, self.ch['type_pie'], width_in=4.5,
                                  caption='Figura 2 — Distribuição por Tipo de Arquivo')
        doc.add_page_break()

        # ── 3. Score por arquivo ──────────────────────────────────────
        self._docx_heading(doc, '3. SCORE DE CRITICIDADE POR ARQUIVO')
        self._docx_body(doc, 'Score calculado: base de severidade + bônus proporcional ao tamanho do arquivo.')
        self._docx_add_chart(doc, self.ch.get('risk_all'), width_in=6.2,
                              caption='Figura 3 — Score de Criticidade por Arquivo')
        self._docx_add_chart(doc, self.ch.get('size'), width_in=6.2,
                              caption='Figura 4 — Volume de Dados Expostos por Arquivo')
        doc.add_page_break()

        # ── 4. Mapa de calor ─────────────────────────────────────────
        self._docx_heading(doc, '4. MAPA DE CALOR DE RISCO MULTIDIMENSIONAL')
        self._docx_body(doc,
            'Avalia cada arquivo em 5 dimensões: Exposição Pública, Volume, Tipo, '
            'Risco de Privacidade e Score Final. Cores quentes = maior risco.')
        self._docx_add_chart(doc, self.ch.get('heatmap'), width_in=6.3,
                              caption='Figura 5 — Mapa de Calor Multidimensional')
        doc.add_page_break()

        # ── 5. Gauges individuais ─────────────────────────────────────
        gauges = self.ch.get('gauges', [])
        if gauges:
            self._docx_heading(doc, '5. ANÁLISE INDIVIDUAL POR ARQUIVO (GAUGES)')
            self._docx_body(doc,
                'Velocímetro de criticidade individual para cada arquivo vulnerável identificado no scan.')
            doc.add_paragraph()

            files_list = self.scan_data.get('files', [])[:len(gauges)]
            max_size   = max((float(f.get('size',0)) for f in files_list), default=1) or 1

            COLS = 3
            for row_start in range(0, len(gauges), COLS):
                chunk = gauges[row_start:row_start+COLS]
                fchunk = files_list[row_start:row_start+COLS]

                # Linha de imagens
                img_tbl = doc.add_table(rows=2, cols=COLS)
                img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for col_i, (gpath, f) in enumerate(zip(chunk, fchunk)):
                    # Imagem
                    ic = img_tbl.rows[0].cells[col_i]
                    ic.width = Inches(2.1)
                    self._docx_remove_borders(ic)
                    pi = ic.paragraphs[0]
                    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if Path(gpath).exists():
                        pi.add_run().add_picture(gpath, width=Inches(2.0))
                    # Caption
                    cc = img_tbl.rows[1].cells[col_i]
                    cc.width = Inches(2.1)
                    self._docx_remove_borders(cc)
                    score = min(100, self.SEV_BASE.get(f.get('severity','LOW'),18) +
                                (float(f.get('size',0)) / max_size) * 12)
                    sev   = f.get('severity','LOW')
                    name  = f.get('key', f.get('name',''))[-20:]
                    pc = cc.paragraphs[0]
                    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = pc.add_run(f"{name}\n{sev} | {score:.0f} pts")
                    rc.font.size = Pt(7.5); rc.font.name = 'Arial'
                    rc.font.color.rgb = RGBColor(0x6b,0x72,0x80)

                # Preencher células vazias se necessário
                for col_i in range(len(chunk), COLS):
                    for row_i in range(2):
                        self._docx_remove_borders(img_tbl.rows[row_i].cells[col_i])

            doc.add_paragraph()
            doc.add_page_break()
            next_sec = 6
        else:
            next_sec = 5

        # ── Vulnerabilidades ──────────────────────────────────────────
        if self.vulnerabilities:
            self._docx_heading(doc, f'{next_sec}. VULNERABILIDADES IDENTIFICADAS')
            for vuln in self.vulnerabilities:
                self._docx_heading(doc, f"{vuln['icon']} {vuln['name']} ({vuln['count']} arquivos)", level=2)
                vt = doc.add_table(rows=1+len(vuln['items'][:10]), cols=4)
                vt.style = 'Table Grid'
                for i, h in enumerate(['Arquivo','Severidade','Tamanho','Motivo']):
                    self._docx_set_bg(vt.rows[0].cells[i], '3b82f6')
                    r = vt.rows[0].cells[i].paragraphs[0].add_run(h)
                    r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
                    r.font.color.rgb=RGBColor(255,255,255)
                for ri, item in enumerate(vuln['items'][:10], 1):
                    vals = [item['file'][-50:], item['severity'],
                            self._format_size(item['size']), item.get('reason','')[:35]]
                    for ci, v in enumerate(vals):
                        c = vt.rows[ri].cells[ci]
                        r = c.paragraphs[0].add_run(v)
                        r.font.size=Pt(8.5); r.font.name='Arial'
                doc.add_paragraph()
            next_sec += 1

        # ── Comparativo ───────────────────────────────────────────────
        comp = self._build_comparison()
        if comp:
            self._docx_heading(doc, f'{next_sec}. COMPARATIVO COM SCAN ANTERIOR')
            self._docx_body(doc, f"Scan anterior: {comp['prev_timestamp']}")
            ct = doc.add_table(rows=5, cols=4)
            ct.style = 'Table Grid'
            for i, h in enumerate(['Métrica','Anterior','Atual','Variação']):
                self._docx_set_bg(ct.rows[0].cells[i], '1e3a8a')
                r = ct.rows[0].cells[i].paragraphs[0].add_run(h)
                r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
                r.font.color.rgb=RGBColor(255,255,255)
            comp_rows = [
                ('Total', comp['prev_total'], comp['curr_total'],
                 comp['curr_total']-comp['prev_total']),
                ('CRÍTICOS',
                 self.previous_scan['severity_distribution'].get('critical',0),
                 sd.get('critical',0), comp['delta_critical']),
                ('ALTOS',
                 self.previous_scan['severity_distribution'].get('high',0),
                 sd.get('high',0), comp['delta_high']),
                ('Risk Score', self.previous_scan.get('risk_score',0),
                 self.scan_data.get('risk_score',0), comp['delta_score']),
            ]
            for ri,(lbl,prev,curr,delta) in enumerate(comp_rows, 1):
                vals = [lbl, str(prev), str(curr), self._delta_str(delta)]
                for ci,v in enumerate(vals):
                    ct.rows[ri].cells[ci].paragraphs[0].add_run(v).font.size = Pt(9)
            next_sec += 1

        # ── Recomendações ─────────────────────────────────────────────
        self._docx_heading(doc, f'{next_sec}. RECOMENDAÇÕES PRIORITÁRIAS')
        effort_map = {'CRÍTICA':'Baixo','ALTA':'Médio','MÉDIA':'Alto','BAIXA':'Baixo'}
        impact_map = {'CRÍTICA':'Máximo','ALTA':'Alto','MÉDIA':'Médio','BAIXA':'Baixo'}
        prio_hex   = {'CRÍTICA':'dc2626','ALTA':'ea580c','MÉDIA':'ca8a04','BAIXA':'16a34a'}
        for rec in self.recommendations:
            # Cabeçalho colorido
            hdr = doc.add_table(rows=1, cols=2)
            hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._docx_set_bg(hdr.rows[0].cells[0], prio_hex.get(rec['priority'],'3b82f6'))
            self._docx_set_bg(hdr.rows[0].cells[1], prio_hex.get(rec['priority'],'3b82f6'))
            hdr.rows[0].cells[0].width = Inches(4.5)
            hdr.rows[0].cells[1].width = Inches(1.9)
            r0 = hdr.rows[0].cells[0].paragraphs[0].add_run(f"[{rec['priority']}]  {rec['title']}")
            r0.bold=True; r0.font.size=Pt(10.5); r0.font.name='Arial'
            r0.font.color.rgb=RGBColor(255,255,255)
            r1 = hdr.rows[0].cells[1].paragraphs[0].add_run(f"Prazo: {rec.get('timeline','')}")
            r1.font.size=Pt(9); r1.font.name='Arial'
            r1.font.color.rgb=RGBColor(255,255,255)
            hdr.rows[0].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

            # Meta
            meta = doc.add_table(rows=1, cols=4)
            meta.style = 'Table Grid'
            for ci,(lbl,val) in enumerate([
                ('Responsável', rec.get('responsible','')),
                ('Esforço',     effort_map.get(rec['priority'],'')),
                ('Impacto',     impact_map.get(rec['priority'],'')),
                ('Prazo',       rec.get('timeline','')),
            ]):
                self._docx_set_bg(meta.rows[0].cells[ci], 'f8fafc')
                p = meta.rows[0].cells[ci].paragraphs[0]
                rl = p.add_run(f"{lbl}: ")
                rl.bold=True; rl.font.size=Pt(9); rl.font.name='Arial'
                rv = p.add_run(val)
                rv.font.size=Pt(9); rv.font.name='Arial'

            # Descrição e ações
            self._docx_body(doc, rec['description'], after=2)
            for a in rec['actions'][:5]:
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                br = bp.add_run(a)
                br.font.size=Pt(9.5); br.font.name='Arial'
            doc.add_paragraph()
        next_sec += 1

        # ── Compliance ────────────────────────────────────────────────
        if self.compliance_status:
            self._docx_heading(doc, f'{next_sec}. STATUS DE CONFORMIDADE')
            ct2 = doc.add_table(rows=1+len(self.compliance_status), cols=3)
            ct2.style = 'Table Grid'
            for i,h in enumerate(['Framework','Status','Observações']):
                self._docx_set_bg(ct2.rows[0].cells[i], '1e3a8a')
                r = ct2.rows[0].cells[i].paragraphs[0].add_run(h)
                r.bold=True; r.font.size=Pt(9); r.font.name='Arial'
                r.font.color.rgb=RGBColor(255,255,255)
            for ri,fw in enumerate(self.compliance_status, 1):
                for ci,v in enumerate([fw['name'],fw['status'],fw['issues']]):
                    ct2.rows[ri].cells[ci].paragraphs[0].add_run(v).font.size = Pt(9)
            next_sec += 1

        # ── Conclusão ─────────────────────────────────────────────────
        self._docx_heading(doc, f'{next_sec}. CONCLUSÕES E PRÓXIMOS PASSOS')
        self._docx_body(doc,
            f"Este relatório identificou {risk['critical_count']} vulnerabilidades críticas "
            f"e {risk['high_count']} de alto risco no ambiente {self.provider_name}. "
            "É fundamental executar as ações prioritárias nos prazos estabelecidos.")
        doc.add_paragraph()
        self._docx_body(doc, "Próximos Passos Recomendados:", bold=True, color='1e3a8a')
        for step in ['Reunião de alinhamento com stakeholders (48h)',
                     'Execução das ações críticas (0-7 dias)',
                     'Implementação de controles preventivos (30-60 dias)',
                     'Re-auditoria de segurança (90 dias)',
                     'Estabelecimento de programa de segurança contínua']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            bp.add_run(step).font.size = Pt(10)

        doc.save(str(filepath))
        report_hash = self._compute_hash(str(filepath))
        (filepath.with_suffix('.sha256')).write_text(
            f"SHA-256: {report_hash}\nArquivo: {filename}\nGerado em: {self.report_date.isoformat()}\n")
        print(f"✅ DOCX: {filepath}")
        return str(filepath)


# ══════════════════════════════════════════════════════════════════════
# FUNÇÃO PÚBLICA (API compatível com versão anterior)
# ══════════════════════════════════════════════════════════════════════
def generate_executive_report(scan_data: dict,
                               client_info: dict = None,
                               output_format: str = 'both') -> dict:
    """
    Gera relatórios PDF e/ou DOCX com gráficos de criticidade.

    Args:
        scan_data:     Dados do scan (bucket, files, severity_distribution, risk_score, provider)
        client_info:   {'name': ..., 'contact': ...}
        output_format: 'pdf' | 'docx' | 'both'

    Returns:
        dict com caminhos gerados: {'pdf': '...', 'docx': '...'}
    """
    print("\n" + "="*60)
    print("🚀 INICIANDO GERAÇÃO DE RELATÓRIO EXECUTIVO")
    print("="*60)
    try:
        gen     = ProfessionalReportGenerator(scan_data, client_info)
        results = {}

        if output_format in ('pdf', 'both'):
            try:
                results['pdf'] = gen.generate_pdf()
            except Exception as e:
                print(f"❌ Erro PDF: {e}")
                import traceback; traceback.print_exc()
                results['pdf_error'] = str(e)

        if output_format in ('docx', 'both') and DOCX_AVAILABLE:
            try:
                results['docx'] = gen.generate_docx()
            except Exception as e:
                print(f"❌ Erro DOCX: {e}")
                import traceback; traceback.print_exc()
                results['docx_error'] = str(e)

        print("\n" + "="*60)
        print("✅ GERAÇÃO CONCLUÍDA")
        print("="*60 + "\n")
        return results

    except Exception as e:
        print(f"\n❌ ERRO CRÍTICO: {e}")
        import traceback; traceback.print_exc()
        raise

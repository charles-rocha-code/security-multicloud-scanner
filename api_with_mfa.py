"""
api_with_mfa.py - FastAPI app do Security Multicloud Storage Scanner com MFA obrigatório.

Inclui:
- Páginas: /login, /mfa/setup, /dashboard
- Auth/MFA:
  POST /auth/register
  POST /auth/login
  GET  /auth/mfa/status
  POST /auth/mfa/verify
  POST /auth/logout
- API protegida (MFA obrigatório):
  GET  /api/dashboard
  GET  /scan/{target}            (scan público - compatível com dashboard)
  POST /scan/public              (scan público - recomendado)
  POST /scan/authenticated       (scan privado com credenciais)
  POST /generate-report          (gera PDF + DOCX no padrão do ambiente)
  GET  /reports_executive/{file} (download público do relatório gerado)

Obs:
- Este arquivo espera que existam no projeto:
  - templates/login.html
  - templates/mfa_setup.html
  - templates/dashboard.html
  - auth_mfa.py
  - auditor_s3_authenticated.py (classe S3AuthenticatedAuditor)
  - auditor_universal.py (opcional, para scan público multicloud)
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from fastapi import Depends, FastAPI, HTTPException, Request , Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, EmailStr, Field

# === Auth/MFA do projeto ===
from auth_mfa import (  # type: ignore
    ACTIVE_SESSIONS,
    USERS_DB,
    get_current_user,
    login_user,
    register_user,
    require_mfa,
    setup_mfa,
    verify_and_activate_mfa,
)

# === Auditor autenticado (AWS S3) ===
from auditor_s3_authenticated import S3AuthenticatedAuditor  # type: ignore

# === PDF/DOCX Profissional ===
try:
    from generate_report import generate_executive_report as _gen_executive_report
    PROFESSIONAL_REPORT = True
    print("✅ Gerador de relatórios profissional carregado")
except ImportError as e:
    PROFESSIONAL_REPORT = False
    print(f"⚠️  generate_report.py não encontrado, usando fallback: {e}")
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas
    from docx import Document


# -----------------------------------------------------------------------------
# App / Templates / Static
# -----------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
STATIC_DIR = BASE_DIR / "static"
REPORTS_DIR = BASE_DIR / "reports_executive"

app = FastAPI(title="Security Multicloud Storage Scanner (MFA)")

# Cache em memória do último scan por usuário (email) - ajuda a gerar relatório
LAST_SCAN_BY_EMAIL: Dict[str, Dict[str, Any]] = {}

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://127.0.0.1:8000",
        "http://localhost:8000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

REPORTS_DIR.mkdir(parents=True, exist_ok=True)
# IMPORTANTE: no seu ambiente AWS, o browser abre diretamente o PDF/DOCX.
# Para evitar 401 no GET do arquivo (porque o browser não manda Authorization),
# exponho o diretório como estático (mesmo comportamento do seu print).
app.mount("/reports_executive", StaticFiles(directory=str(REPORTS_DIR)), name="reports_executive")


# -----------------------------------------------------------------------------
# Schemas
# -----------------------------------------------------------------------------
class UserRegisterIn(BaseModel):
    email: EmailStr
    full_name: str = Field(min_length=1, max_length=120)
    password: str = Field(min_length=4, max_length=256)


class UserLoginIn(BaseModel):
    email: EmailStr
    password: str
    mfa_code: Optional[str] = None


class MFASetupIn(BaseModel):
    email: EmailStr


class MFAVerifyIn(BaseModel):
    email: EmailStr
    code: str = Field(min_length=6, max_length=10)


@app.post("/auth/mfa/setup")
def auth_mfa_setup(data: MFASetupIn):
    """Gera QR Code para MFA após o login"""
    try:
        email = str(data.email).lower().strip()
        
        # Validar se usuário existe
        user = USERS_DB.get(email)
        if not user:
            raise HTTPException(status_code=404, detail="Usuário não encontrado")
        
        # ✅ CORRIGIDO: Chamada simplificada sem senha
        mfa = setup_mfa(email=email)
        
        return {
            "email": email,
            "qr_code": mfa.get("qr_code"),
            "secret": mfa.get("secret"),
            "backup_codes": mfa.get("backup_codes", []),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


class PublicScanIn(BaseModel):
    target: str = Field(..., description="Bucket/host alvo. Ex.: bucket.s3.amazonaws.com, storage.googleapis.com/bucket, *.blob.core.windows.net")
    max_objects: int = 1000


class AuthenticatedScanIn(BaseModel):
    provider: str = Field(default="AWS_S3")
    bucket: str

    # AWS
    aws_access_key_id: Optional[str] = None
    aws_secret_access_key: Optional[str] = None
    aws_session_token: Optional[str] = None
    region_name: Optional[str] = None  # nome correto (o dashboard às vezes manda "region")

    # compat: alguns dashboards mandam "region"
    region: Optional[str] = None
    # GCS
    service_account_key: Optional[Dict[str, Any]] = None


    max_objects: int = 1000


class GenerateReportIn(BaseModel):
    scan_result: Dict[str, Any] = Field(default_factory=dict)
    report_title: str = "Relatório Executivo de Segurança"
    client_name: Optional[str] = None


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------
def _safe_filename(name: str) -> str:
    keep = []
    for ch in str(name):
        if ch.isalnum() or ch in ("-", "_", ".", "@"):
            keep.append(ch)
        else:
            keep.append("_")
    return "".join(keep)


def _format_bytes(num: Any) -> str:
    try:
        n = float(num or 0)
    except Exception:
        return str(num)
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if n < 1024.0:
            return f"{n:.2f} {unit}"
        n /= 1024.0
    return f"{n:.2f} PB"

import re

def _detect_provider(target: str) -> str:
    t = (target or "").lower().strip()

    # remove esquema
    if t.startswith("http://"):
        t = t[len("http://"):]
    if t.startswith("https://"):
        t = t[len("https://"):]

    # AWS S3 (pega s3-REGION e s3.REGION e s3.amazonaws.com)
    if (
        re.search(r"\.s3([.-][a-z0-9-]+)?\.amazonaws\.com($|/)", t)
        or re.search(r"\.s3\.[a-z0-9-]+\.amazonaws\.com($|/)", t)
    ):
        return "AWS_S3"

    # GCS
    if "storage.googleapis.com" in t or ".storage.googleapis.com" in t:
        return "GCS"

    # Azure Blob
    if ".blob.core.windows.net" in t:
        return "AZURE_BLOB"

    return "UNIVERSAL"



def _load_universal_auditor():
    """
    Tenta carregar um auditor multicloud do seu projeto (auditor_universal.py).
    Não falha se não existir.
    """
    try:
        import auditor_universal  # type: ignore

        for name in ("UniversalAuditor", "AuditorUniversal", "UniversalScanner", "UniversalStorageAuditor"):
            cls = getattr(auditor_universal, name, None)
            if cls:
                return cls
    except Exception:
        return None
    return None



def _ensure_recommendations(scan: Dict[str, Any], provider: str) -> None:
    """Garante recomendações mínimas quando não houver achados (para AWS/GCP/Azure)."""
    recos = scan.get('recommendations')
    if recos is None:
        recos = []
    if not isinstance(recos, list):
        recos = []

    files = scan.get('files') or scan.get('vulnerable_files') or []
    has_findings = isinstance(files, list) and len(files) > 0

    # tenta inferir risco
    risk_counts = scan.get('risk_counts') or {}
    if isinstance(risk_counts, dict):
        total_risks = sum(int(v or 0) for v in risk_counts.values() if str(v).isdigit() or isinstance(v,(int,float)))
    else:
        total_risks = 0

    if has_findings or total_risks > 0 or recos:
        scan['recommendations'] = recos
        return

    p = (provider or scan.get('provider') or 'UNIVERSAL').upper()
    base = [
        'Nenhum risco crítico detectado no momento.',
        'Mantenha boas práticas de segurança e monitoramento contínuo.',
    ]
    if p in ('AWS', 'AWS_S3', 'S3'):
        base += [
            'Revisar Block Public Access, ACLs e Bucket Policy do S3.',
            'Ativar CloudTrail/CloudWatch e alertas para mudanças de permissão.',
        ]
    elif p in ('GCP', 'GCS', 'GOOGLE', 'GOOGLE_CLOUD_STORAGE'):
        base += [
            'Revisar IAM do bucket e habilitar Public Access Prevention (quando aplicável).',
            'Habilitar logs de acesso e alertas para alterações de permissões.',
        ]
    elif p in ('AZURE', 'AZURE_BLOB', 'BLOB', 'AZURE_STORAGE'):
        base += [
            'Revisar o nível de acesso público do container (Private/Blob/Container).',
            'Ativar Microsoft Defender for Storage e logs/alertas de acesso.',
        ]
    else:
        base += [
            'Revisar permissões, políticas de acesso público e logs do storage.',
        ]

    scan['recommendations'] = base
def _normalize_scan_result(scan: Dict[str, Any], provider: str, target: str) -> Dict[str, Any]:
    # garante campos base usados no dashboard/relatório
    scan = scan or {}
    scan.setdefault("provider", provider)
    scan.setdefault("bucket", target)
    scan.setdefault("region", scan.get("region") or scan.get("region_name") or "-")
    scan.setdefault("generated_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    _ensure_recommendations(scan, provider)
    return scan


def _top_findings(scan: Dict[str, Any], limit: int = 12) -> list[str]:
    files = scan.get("files") or scan.get("vulnerable_files") or []
    if not isinstance(files, list):
        return []
    order = {"CRITICAL": 0, "Critical": 0, "HIGH": 1, "High": 1, "MEDIUM": 2, "Medium": 2, "LOW": 3, "Low": 3}

    def key_fn(f: Dict[str, Any]) -> Tuple[int, int]:
        sev = f.get("severity", "LOW")
        size = f.get("size", 0) or 0
        try:
            size_i = int(size)
        except Exception:
            size_i = 0
        return (order.get(str(sev), 9), -size_i)

    files_sorted = sorted(files, key=key_fn)[:limit]
    out: list[str] = []
    for f in files_sorted:
        key = f.get("key") or f.get("path") or f.get("name") or "-"
        sev = f.get("severity", "-")
        size = _format_bytes(f.get("size", 0))
        out.append(f"{sev}: {key} ({size})")
    return out


def _make_pdf(scan: Dict[str, Any], title: str, client_name: Optional[str]) -> Path:
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    bucket = _safe_filename(scan.get("bucket", "target"))
    filename = f"relatorio_{bucket}_{ts}.pdf"
    path = REPORTS_DIR / filename

    c = rl_canvas.Canvas(str(path), pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 18)
    c.drawString(40, height - 60, title)

    c.setFont("Helvetica", 11)
    c.drawString(40, height - 82, f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    if client_name:
        c.drawString(40, height - 100, f"Cliente: {client_name}")

    provider = scan.get("provider", "-")
    target = scan.get("bucket", "-")
    region = scan.get("region", "-")
    c.drawString(40, height - 122, f"Provider: {provider} | Alvo: {target} | Região: {region}")

    y = height - 160
    c.setFont("Helvetica-Bold", 13)
    c.drawString(40, y, "Resumo")
    y -= 18

    summary = scan.get("summary") or {}
    objects_scanned = summary.get("objects_scanned", summary.get("total_files", scan.get("total_files", 0)))
    total_size = summary.get("total_size_bytes", summary.get("total_size", scan.get("total_size_bytes", 0)))
    public_access = scan.get("public_access", scan.get("is_public", False))
    risk_score = scan.get("risk_score", scan.get("risk_total", scan.get("risk", "-")))

    c.setFont("Helvetica", 11)
    c.drawString(40, y, f"- Objetos analisados: {objects_scanned}")
    y -= 16
    c.drawString(40, y, f"- Tamanho total analisado: {_format_bytes(total_size)}")
    y -= 16
    c.drawString(40, y, f"- Exposição: {'Público' if public_access else 'Privado'}")
    y -= 16
    c.drawString(40, y, f"- Risco (consolidado): {risk_score}")
    y -= 22

    c.setFont("Helvetica-Bold", 13)
    c.drawString(40, y, "Achados (amostra)")
    y -= 18
    c.setFont("Helvetica", 10)

    for line in _top_findings(scan, limit=14):
        if y < 80:
            c.showPage()
            y = height - 60
            c.setFont("Helvetica", 10)
        c.drawString(45, y, f"• {line[:145]}")
        y -= 14

    recs = scan.get("recommendations") or []
    if isinstance(recs, list) and recs:
        y -= 8
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, "Recomendações (top)")
        y -= 18
        c.setFont("Helvetica", 10)
        for rec in recs[:10]:
            if y < 80:
                c.showPage()
                y = height - 60
                c.setFont("Helvetica", 10)
            c.drawString(45, y, f"• {str(rec)[:145]}")
            y -= 14

    c.showPage()
    c.save()
    return path


def _make_docx(scan: Dict[str, Any], title: str, client_name: Optional[str]) -> Path:
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    bucket = _safe_filename(scan.get("bucket", "target"))
    filename = f"relatorio_{bucket}_{ts}.docx"
    path = REPORTS_DIR / filename

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    if client_name:
        doc.add_paragraph(f"Cliente: {client_name}")

    provider = scan.get("provider", "-")
    target = scan.get("bucket", "-")
    region = scan.get("region", "-")
    doc.add_paragraph(f"Provider: {provider} | Alvo: {target} | Região: {region}")

    doc.add_heading("Resumo", level=1)
    summary = scan.get("summary") or {}
    objects_scanned = summary.get("objects_scanned", summary.get("total_files", scan.get("total_files", 0)))
    total_size = summary.get("total_size_bytes", summary.get("total_size", scan.get("total_size_bytes", 0)))
    public_access = scan.get("public_access", scan.get("is_public", False))
    risk_score = scan.get("risk_score", scan.get("risk_total", scan.get("risk", "-")))

    doc.add_paragraph(f"Objetos analisados: {objects_scanned}")
    doc.add_paragraph(f"Tamanho total analisado: {_format_bytes(total_size)}")
    doc.add_paragraph(f"Exposição: {'Público' if public_access else 'Privado'}")
    doc.add_paragraph(f"Risco (consolidado): {risk_score}")

    doc.add_heading("Achados (amostra)", level=1)
    for line in _top_findings(scan, limit=20):
        doc.add_paragraph(line, style="List Bullet")

    recs = scan.get("recommendations") or []
    if isinstance(recs, list) and recs:
        doc.add_heading("Recomendações (top)", level=1)
        for rec in recs[:15]:
            doc.add_paragraph(str(rec), style="List Bullet")

    doc.save(str(path))
    return path


# -----------------------------------------------------------------------------
# Pages
# -----------------------------------------------------------------------------
@app.get("/", include_in_schema=False)
def root():
    return RedirectResponse(url="/login")


@app.get("/login", response_class=HTMLResponse, include_in_schema=False)
def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/mfa/setup", response_class=HTMLResponse, include_in_schema=False)
def mfa_setup_page(request: Request):
    return templates.TemplateResponse("mfa_setup.html", {"request": request})


@app.get("/dashboard", response_class=HTMLResponse, include_in_schema=False)
def dashboard_page(request: Request):
    # Proteção real é via /api/dashboard (com Authorization)
    return templates.TemplateResponse("dashboard.html", {"request": request})


# -----------------------------------------------------------------------------
# Health
# -----------------------------------------------------------------------------
@app.get("/health")
def health():
    return {"status": "ok"}


# -----------------------------------------------------------------------------
# Auth/MFA API
# -----------------------------------------------------------------------------
@app.post("/auth/register")
def auth_register(data: UserRegisterIn):
    register_user(email=data.email, password=data.password, full_name=data.full_name)
    # ✅ CORRIGIDO: setup_mfa agora só precisa do email
    mfa = setup_mfa(email=data.email)
    return {
        "registered": True,
        "mfa_setup_required": True,
        "qr_code": mfa["qr_code"],
        "backup_codes": mfa.get("backup_codes", []),
        "email": data.email,
    }


@app.post("/auth/login")
def auth_login(data: UserLoginIn, response: Response):
    return login_user(email=data.email, password=data.password, mfa_code=data.mfa_code, response=response)


@app.get("/auth/mfa/status")
def auth_mfa_status(email: EmailStr):
    user = USERS_DB.get(str(email))
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    return {
        "email": user.get("email"),
        "mfa_enabled": bool(user.get("mfa_enabled")),
        "has_secret": bool(user.get("mfa_secret")),
    }


@app.post("/auth/mfa/verify")
def auth_mfa_verify(data: MFAVerifyIn):
    return verify_and_activate_mfa(email=data.email, code=data.code)


@app.post("/auth/logout")
def auth_logout(user=Depends(get_current_user)):
    email = user.get("email")
    removed = 0
    to_delete = [t for t, s in ACTIVE_SESSIONS.items() if s.get("email") == email]
    for t in to_delete:
        del ACTIVE_SESSIONS[t]
        removed += 1
    return {"ok": True, "removed_sessions": removed}


# -----------------------------------------------------------------------------
# API protegida (MFA obrigatório)
# -----------------------------------------------------------------------------
@app.get("/api/dashboard")
def dashboard_api(user=Depends(require_mfa)):
    return {"ok": True, "user": user}


# -----------------------------------------------------------------------------
# Scan público (compatível com dashboard)
# -----------------------------------------------------------------------------
@app.get("/scan/{target:path}")
def scan_public_compat(target: str, user=Depends(require_mfa)):
    payload = PublicScanIn(target=target)
    return scan_public(payload, user)


@app.post("/scan/public")
def scan_public(payload: PublicScanIn, user=Depends(require_mfa)):
    provider = _detect_provider(payload.target)

    # Fast-path: para endpoints/buckets AWS S3, use o auditor S3 nativo (evita inconsistências de assinatura do universal).
    if provider == 'AWS_S3':
        try:
            from auditor import S3Auditor  # type: ignore
            aws_aud = S3Auditor(payload.target, max_objects=payload.max_objects)
            result = aws_aud.run()
            result = _normalize_scan_result(result, provider=provider, target=payload.target)
            email = user.get('email') if isinstance(user, dict) else None
            if email:
                LAST_SCAN_BY_EMAIL[email] = result
            return result
        except Exception:
            # se algo falhar, cai no fluxo universal abaixo
            pass

    Universal = _load_universal_auditor()
    if not Universal:
        raise HTTPException(
            status_code=501,
            detail="auditor_universal.py não encontrado (scan público multicloud indisponível neste ambiente).",
        )

    # Compat com diferentes assinaturas do UniversalAuditor.
    # Queremos SEMPRE repassar max_objects quando possível.
    auditor = None
    for ctor_kwargs in (
        {"target": payload.target, "max_objects": payload.max_objects, "region": "global"},
        {"target": payload.target, "max_objects": payload.max_objects},
        {"domain": payload.target, "max_objects": payload.max_objects},
        {"domain": payload.target},
    ):
        try:
            auditor = Universal(**ctor_kwargs)  # type: ignore
            break
        except TypeError:
            continue

    if auditor is None:
        # fallback final: positional
        try:
            auditor = Universal(payload.target, payload.max_objects)  # type: ignore
        except Exception:
            auditor = Universal(payload.target)  # type: ignore

    result = auditor.run() if hasattr(auditor, "run") else auditor.scan()  # type: ignore
    result = _normalize_scan_result(result, provider=provider, target=payload.target)

    # cache para relatório
    email = user.get("email") if isinstance(user, dict) else None
    if email:
        LAST_SCAN_BY_EMAIL[email] = result

    return result


# -----------------------------------------------------------------------------
# Scan autenticado (AWS S3)
# -----------------------------------------------------------------------------
@app.post("/scan/authenticated")
def scan_authenticated(payload: AuthenticatedScanIn, user=Depends(require_mfa)):
    provider = (payload.provider or "AWS_S3").upper()

    # GCS AUTENTICADO
    if provider == "GCS":
        from auditor_gcs_authenticated import GCSAuthenticatedAuditor
        if not payload.service_account_key:
            raise HTTPException(status_code=400, detail="Credencial GCS ausente (service_account_key JSON).")
        auditor = GCSAuthenticatedAuditor(
            bucket_name=payload.bucket,
            service_account_key=payload.service_account_key,
            max_objects=payload.max_objects,
        )
        result = auditor.run()
        result = _normalize_scan_result(result, provider="GCS", target=payload.bucket)

    # AWS S3 AUTENTICADO
    elif provider == "AWS_S3":
        region_name = payload.region_name or payload.region
        if not payload.aws_access_key_id or not payload.aws_secret_access_key:
            raise HTTPException(status_code=400, detail="Credenciais AWS ausentes (aws_access_key_id / aws_secret_access_key).")
        auditor = S3AuthenticatedAuditor(
            bucket_name=payload.bucket,
            aws_access_key_id=payload.aws_access_key_id,
            aws_secret_access_key=payload.aws_secret_access_key,
            aws_session_token=payload.aws_session_token,
            region_name=region_name,
            max_objects=payload.max_objects,
        )
        result = auditor.run()
        result = _normalize_scan_result(result, provider="AWS_S3", target=payload.bucket)

    else:
        raise HTTPException(status_code=400, detail=f"Provider '{provider}' não suportado.")

    email = user.get("email") if isinstance(user, dict) else None
    if email:
        LAST_SCAN_BY_EMAIL[email] = result

    return result


# Compatibilidade com versões antigas do dashboard
@app.post("/scan/airdrop")
def scan_airdrop(payload: AuthenticatedScanIn, user=Depends(require_mfa)):
    return scan_authenticated(payload, user)


@app.post("/scan/amazetest")
def scan_amazetest(payload: AuthenticatedScanIn, user=Depends(require_mfa)):
    return scan_authenticated(payload, user)


# -----------------------------------------------------------------------------
# Relatório Executivo (PDF + DOCX) no padrão do ambiente AWS
# -----------------------------------------------------------------------------
@app.post("/generate-report")
def generate_report(data: GenerateReportIn, user=Depends(require_mfa)):
    scan: Dict[str, Any] = data.scan_result or {}
    email = user.get("email") if isinstance(user, dict) else None

    if not scan and email and email in LAST_SCAN_BY_EMAIL:
        scan = LAST_SCAN_BY_EMAIL[email]

    if not scan:
        raise HTTPException(status_code=400, detail="Nenhum scan disponível para gerar relatório. Execute um scan antes.")

    # Normaliza pra evitar relatório vazio
    provider = scan.get("provider") or _detect_provider(str(scan.get("bucket", "")))
    target = str(scan.get("bucket") or scan.get("target") or "-")
    scan = _normalize_scan_result(scan, provider=provider, target=target)

    # ── Usar gerador profissional com gráficos ────────────────────────
    if PROFESSIONAL_REPORT:
        try:
            client_info = {
                "name":    data.client_name or "Cliente",
                "contact": email or "-",
            }
            results = _gen_executive_report(scan, client_info=client_info, output_format="both")
            pdf_file  = Path(results.get("pdf",  ""))
            docx_file = Path(results.get("docx", ""))
            return {
                "success": True,
                "files": {
                    "pdf":  f"reports_executive/{pdf_file.name}"  if pdf_file.exists()  else None,
                    "docx": f"reports_executive/{docx_file.name}" if docx_file.exists() else None,
                },
                "message": "Relatórios profissionais gerados com sucesso",
            }
        except Exception as e:
            print(f"⚠️  Erro no gerador profissional, usando fallback: {e}")

    # ── Fallback simples ──────────────────────────────────────────────
    pdf_path  = _make_pdf(scan, data.report_title, data.client_name)
    docx_path = _make_docx(scan, data.report_title, data.client_name)
    return {
        "success": True,
        "files": {
            "pdf":  f"reports_executive/{pdf_path.name}",
            "docx": f"reports_executive/{docx_path.name}",
        },
        "message": "Relatórios gerados com sucesso",
    }


# Rota opcional (quando quiser forçar download com filename)
@app.get("/download-report/{filename}")
def download_report(filename: str, user=Depends(require_mfa)):
    safe = _safe_filename(filename)
    path = REPORTS_DIR / safe
    if not path.exists():
        raise HTTPException(status_code=404, detail="Relatório não encontrado")
    media = "application/pdf" if safe.lower().endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(str(path), media_type=media, filename=safe)
# ==================================================
# EXECUÇÃO DIRETA DO SERVIDOR (quando rodar python)
# ==================================================
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api_with_mfa:app", host="0.0.0.0", port=8000, reload=False)